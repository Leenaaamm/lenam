from docx import Document
from docx.shared import Pt, Inches
import sys

def hidden_binary_string(doc_path, binary_string):
    document = Document(doc_path)
    for paragraph in document.paragraphs:
        index = 0
        for run in paragraph.runs:
            text = run.text
            for char in text:
                new_run = paragraph.add_run(char)
                # Sao chép định dạng từ run cũ
                new_run.bold = run.bold
                new_run.italic = run.italic
                new_run.underline = run.underline
                new_run.font.name = run.font.name
                if index >= len(binary_string   ):
                    continue
                if char == " " :
                    if binary_string[index] == "1":
                        new_run.font.size = Pt(run.font.size.pt + 1)
                    elif binary_string[index] == "0":
                        new_run.font.size = Pt(run.font.size.pt - 1)
                    index += 1
                else:
                    new_run.font.size = run.font.size

            # Xóa run cũ sau khi đã tạo run mới từ từng ký tự
            run.text = ""

    document.save("hidden.docx")
    print("Đã lưu file 'hidden.docx' với các thuộc tính đã chỉnh sửa.")

if __name__ == "__main__":
    file_directory = sys.argv[1]
    binary_path = sys.argv[2]
    with open(binary_path, "r") as binary_file:
        binary_string = binary_file.read().strip()
    if file_directory.endswith(".docx") is False:
        print("Lỗi: Đường dẫn không phải là file .docx")
        sys.exit(1)

    hidden_binary_string(file_directory,binary_string)
